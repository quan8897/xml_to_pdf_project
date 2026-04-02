"""
Tạo file template docx với placeholder Jinja2 từ file gốc.
Chạy 1 lần để tạo: mau-01-tts-template.docx
"""
import docx
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import re

src = 'mau-01-tts-to-khai-doi-voi-hoat-dong-cho-thue-tai-san.docx'
dst = 'mau-01-tts-template.docx'

doc = Document(src)

# Map: chuỗi cần thay → placeholder Jinja2
# Dùng re.sub để thay thế
REPLACEMENTS = [
    # Checkbox loại khai
    ('cá nhân ủy quyền theo quy định của pháp luật dân sự □',
     'cá nhân ủy quyền theo quy định của pháp luật dân sự {{ chk_ds }}'),
    ('nộp thuê thay theo pháp luật thuế □',
     'nộp thuê thay theo pháp luật thuế {{ chk_thue }}'),

    # [01] Kỳ tính thuế
    ('[01a] Năm ...', '[01a] Năm {{ ky_nam }}'),
    ('từ ngày ... tháng ... năm ... đến ngày ngày ... tháng ... năm ...',
     'từ ngày {{ ky_tu_ngay }} đến ngày {{ ky_den_ngay }}'),
    ('[01c] Tháng ... năm ...', '[01c] Tháng {{ ky_thang }} năm {{ ky_nam }}'),
    ('[01d] Quý ... năm ... (Từ tháng .../... đến tháng .../...)',
     '[01d] Quý {{ ky_quy }} năm {{ ky_nam }}'),

    # [02][03] lần đầu / bổ sung
    ('[02] Lần đầu: □\t[03] Bổ sung lần thứ: ...',
     '[02] Lần đầu: {{ chk_lan_dau }}\t[03] Bổ sung lần thứ: {{ so_lan }}'),

    # [04]-[09]
    ('[04] Người nộp thuế:\t ', '[04] Người nộp thuế:\t{{ ten_nnt }}'),
    ('[06] Địa chỉ liên hệ:\t', '[06] Địa chỉ liên hệ:\t{{ dchi }}'),
    ('[07] Điện thoại:\t[08] Fax:\t[09] Email:\t',
     '[07] Điện thoại:\t{{ dthoai }}\t[08] Fax:\t{{ fax }}\t[09] Email:\t{{ email }}'),
    ('[10] Số CMND (trường hợp cá nhân quốc tịch Việt Nam):\t',
     '[10] Số CMND (trường hợp cá nhân quốc tịch Việt Nam):\t{{ ct10 }}'),
    ('[11] Hộ chiếu (trường hợp cá nhân không có quốc tịch Việt nam):\t',
     '[11] Hộ chiếu (trường hợp cá nhân không có quốc tịch Việt nam):\t{{ ct11 }}'),

    # [12a][12b]
    ('[12a] Ngày sinh:\t/\t/\t [12b]\tQuốc tịch:\t',
     '[12a] Ngày sinh:\t{{ ct12a }}\t [12b]\tQuốc tịch:\t{{ ct12b }}'),
    ('[12c] Số CMND/CCCD:\t[12c.1] Ngày cấp:\t[12c.2]\tNơi cấp:\t',
     '[12c] Số CMND/CCCD:\t{{ ct12c_so }}\t[12c.1] Ngày cấp:\t{{ ct12c_ngay }}\t[12c.2]\tNơi cấp:\t{{ ct12c_noi }}'),

    # [12d][12đ][12e][12f]
    ('[12d] Số hộ chiếu:\t ………… [12d.1] Ngày cấp: ....... [12d.2] Nơi cấp:\t',
     '[12d] Số hộ chiếu:\t{{ ct12d_so }} [12d.1] Ngày cấp: {{ ct12d_ngay }} [12d.2] Nơi cấp:\t{{ ct12d_noi }}'),
    ('[12đ] Số giấy thông hành (đối với thương nhân nước ngoài):\t',
     '[12đ] Số giấy thông hành (đối với thương nhân nước ngoài):\t{{ ct12dd_so }}'),
    ('[12đ.1] Ngày cấp:\t[12đ.2]\tNơi cấp:\t',
     '[12đ.1] Ngày cấp:\t{{ ct12dd_ngay }}\t[12đ.2]\tNơi cấp:\t{{ ct12dd_noi }}'),
    ('[12e] Số CMND biên giới (đối với thương nhân nước ngoài):\t',
     '[12e] Số CMND biên giới (đối với thương nhân nước ngoài):\t{{ ct12e_so }}'),
    ('[12e.1] Ngày cấp:\t[12e.2]\tNơi cấp:\t',
     '[12e.1] Ngày cấp:\t{{ ct12e_ngay }}\t[12e.2]\tNơi cấp:\t{{ ct12e_noi }}'),
    ('[12f] Số Giấy tờ chứng thực cá nhân khác:\t',
     '[12f] Số Giấy tờ chứng thực cá nhân khác:\t{{ ct12f_so }}'),
    ('[12f.1] Ngày cấp:\t[12f.2]\tNơi cấp:\t ',
     '[12f.1] Ngày cấp:\t{{ ct12f_ngay }}\t[12f.2]\tNơi cấp:\t{{ ct12f_noi }}'),

    # [12g] địa chỉ thường trú
    ('[12g.1] Số nhà, đường phố/xóm/ấp/thôn:\t',
     '[12g.1] Số nhà, đường phố/xóm/ấp/thôn:\t{{ ct12g_1 }}'),
    ('[12g.2] Phường/xã/Thị trấn:\t',
     '[12g.2] Phường/xã/Thị trấn:\t{{ ct12g_2 }}'),
    ('[12g.3] Quận/Huyện/Thị xã/Thành phố thuộc tỉnh:\t',
     '[12g.3] Quận/Huyện/Thị xã/Thành phố thuộc tỉnh:\t{{ ct12g_3 }}'),
    ('[12g.4] Tỉnh/Thành phố:\t',
     '[12g.4] Tỉnh/Thành phố:\t{{ ct12g_4 }}'),

    # [12h] chỗ ở hiện tại
    ('[12h] Chỗ ở hiện tại:\t',
     '[12h] Chỗ ở hiện tại:\t{{ ct12h }}'),
    ('[12h.1] Số nhà, đường phố/xóm/ấp/thôn:\t',
     '[12h.1] Số nhà, đường phố/xóm/ấp/thôn:\t{{ ct12h_1 }}'),
    ('[12h.2] Phường/xã/Thị trấn:\t',
     '[12h.2] Phường/xã/Thị trấn:\t{{ ct12h_2 }}'),
    ('[12h.3] Quận/Huyện/Thị xã/Thành phố thuộc tỉnh:\t',
     '[12h.3] Quận/Huyện/Thị xã/Thành phố thuộc tỉnh:\t{{ ct12h_3 }}'),
    ('[12h.4] Tỉnh/Thành phố:\t',
     '[12h.4] Tỉnh/Thành phố:\t{{ ct12h_4 }}'),

    # [12i][12k]
    ('[12i] Giấy chứng nhận đăng ký hộ kinh doanh (nếu có): số:\t',
     '[12i] Giấy chứng nhận đăng ký hộ kinh doanh (nếu có): số:\t{{ ct12i_so }}'),
    ('[12i.1] Ngày cấp: .............. [12i.2] Cơ quan cấp:\t',
     '[12i.1] Ngày cấp: {{ ct12i_ngay }} [12i.2] Cơ quan cấp:\t{{ ct12i_cq }}'),
    ('[12k] Vốn kinh doanh (đồng):\t',
     '[12k] Vốn kinh doanh (đồng):\t{{ ct12k }}'),

    # [13][14][15] đại lý thuế
    ('[13] Tên đại lý thuế (nếu có):\t',
     '[13] Tên đại lý thuế (nếu có):\t{{ ct13 }}'),
    ('[15] Hợp đồng đại lý thuế: số ….. ngày …/…/…..',
     '[15] Hợp đồng đại lý thuế: số {{ ct15_so }} ngày {{ ct15_ngay }}'),

    # [16-22]
    ('[16] Tổ chức khai, nộp thuế thay (nếu có): …………………',
     '[16] Tổ chức khai, nộp thuế thay (nếu có): {{ tc16 }}'),
    ('[18] Địa chỉ:\t', '[18] Địa chỉ:\t{{ tc18 }}'),
    ('[19] Điện thoại:\t [20]\tFax: ..... [21] Email: ......',
     '[19] Điện thoại:\t{{ tc19 }} [20]\tFax: {{ tc20 }} [21] Email: {{ tc21 }}'),
    ('[22] Văn bản ủy quyền (nếu có): số\tngày… tháng...năm ....',
     '[22] Văn bản uỷ quyền (nếu có): số {{ ct22_so }} ngày {{ ct22_ngay }}'),

    # Ngày ký
    ('..., ngày ... tháng ... năm ...', '..., ngày {{ ngay_ky }}'),
]

def replace_in_para(para, old, new):
    """Thay thế text trong paragraph, giữ nguyên formatting của run đầu."""
    full_text = para.text
    if old not in full_text:
        return False
    # Xây lại paragraph với text mới
    new_text = full_text.replace(old, new)
    # Xóa tất cả runs cũ
    for run in para.runs:
        run.text = ''
    # Đặt text mới vào run đầu tiên (hoặc tạo mới)
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)
    return True

# Thực hiện thay thế
changed = 0
for para in doc.paragraphs:
    for old, new in REPLACEMENTS:
        if old in para.text:
            replace_in_para(para, old, new)
            changed += 1
            print(f'  ✓ Replaced: {old[:50]}...')

# Thay thế trong tables (Table 1,2,3 = MST boxes, Table 4 = số liệu, Table 5 = ký tên)
# MST [05]: Table 1
# MST [14]: Table 2  
# MST [17]: Table 3

# Table 4: điền placeholder vào cột Số tiền
codes = ['ct23', 'ct24', 'ct25', 'ct26', 'ct27', 'ct28', 'ct29']
for ri, row in enumerate(doc.tables[4].rows):
    if ri == 0: continue  # skip header
    idx = ri - 1
    if idx < len(codes):
        last_cell = row.cells[-1]
        ph = '{{ ' + codes[idx] + ' }}'
        # Clear và set text
        for p in last_cell.paragraphs:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = ph
            else:
                p.add_run(ph)
        if not last_cell.paragraphs:
            last_cell.add_paragraph(ph)

doc.save(dst)
print(f'\nDone! {changed} replacements. Saved to: {dst}')
