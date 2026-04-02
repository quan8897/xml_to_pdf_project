import docx

doc = docx.Document('mau-01-tts-to-khai-doi-voi-hoat-dong-cho-thue-tai-san.docx')

print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[P{i:03d}] style={p.style.name!r} | {p.text[:120]}')

print()
print('=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---')
    for ri, row in enumerate(table.rows):
        cells = []
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip().replace('\n', ' ')[:60]
            cells.append('[' + txt + ']')
        print('  R' + str(ri) + ': ' + '  '.join(cells))
    print()
